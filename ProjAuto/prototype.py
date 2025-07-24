import streamlit as st
from docx import Document
from docx.shared import Inches # 이미지를 삽입할 경우 사용될 수 있습니다. 현재 코드에서는 직접 사용되지 않음.
import io
import pandas as pd
from datetime import datetime

# --- 1. ID 생성 로직 ---
def generate_requirement_id(category):
    """
    요구사항 분류에 따라 ID를 자동으로 생성합니다.
    예: RQ-F-001, RQ-NF-001
    """
    # Streamlit의 session_state를 사용하여 ID 카운터를 세션 내에서 유지합니다.
    if 'next_id_f' not in st.session_state:
        st.session_state.next_id_f = 1
    if 'next_id_nf' not in st.session_state:
        st.session_state.next_id_nf = 1

    if category == "기능":
        current_id = st.session_state.next_id_f
        st.session_state.next_id_f += 1
        return f"RQ-F-{current_id:03d}"
    elif category == "비기능":
        current_id = st.session_state.next_id_nf
        st.session_state.next_id_nf += 1
        return f"RQ-NF-{current_id:03d}"
    else:
        return "N/A" # 분류가 불확실한 경우

# --- 2. 문서 생성 로직 ---
def create_document(project_info, requirements_df):
    """
    입력받은 프로젝트 정보와 요구사항 데이터를 기반으로 .docx 문서를 생성합니다.
    이 프로토타입에서는 별도의 템플릿 파일 없이 코드 내에서 기본적인 문서 구조를 생성합니다.
    """
    document = Document()

    # 프로젝트 기본 정보 추가
    document.add_heading(f'{project_info["project_title"]} 프로젝트 설계서', level=0)
    document.add_paragraph(f'작성자: {project_info["author"]}')
    document.add_paragraph(f'작성일: {project_info["date"]}')
    document.add_paragraph(f'요약: {project_info["summary"]}')
    document.add_page_break() # 다음 섹션을 새 페이지에서 시작

    # 프로젝트 개요 섹션 (더미 텍스트)
    document.add_heading('1. 프로젝트 개요', level=1)
    document.add_paragraph('본 문서는 ' + project_info["project_title"] + ' 프로젝트의 설계서입니다. 프로젝트의 목표와 개요, 그리고 기대 효과에 대한 내용을 포함합니다. 본 프로토타입은 문서 자동화 기능을 시연하기 위해 제작되었습니다.')
    
    # 요구사항 목록 섹션
    document.add_heading('2. 요구사항 목록', level=1)

    if not requirements_df.empty:
        # 요구사항 테이블 추가
        # 컬럼명: '요구사항ID', '분류', '요구사항 명', '요구사항 설명'
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid' # 표 테두리 스타일 적용 (Word 기본 스타일)
        
        # 헤더 행 추가
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '요구사항ID'
        hdr_cells[1].text = '분류'
        hdr_cells[2].text = '요구사항 명'
        hdr_cells[3].text = '요구사항 설명'

        # 데이터 행 추가
        for index, row in requirements_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['요구사항ID'])
            row_cells[1].text = str(row['분류'])
            row_cells[2].text = str(row['요구사항 명'])
            row_cells[3].text = str(row['요구사항 설명'])
    else:
        document.add_paragraph("아직 등록된 요구사항이 없습니다.")

    # 문서를 메모리 내 바이너리 스트림으로 저장
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0) # 스트림의 시작 지점으로 커서 이동 (읽기 위해 필수)
    return buffer

# --- 3. Streamlit UI 구성 ---
st.set_page_config(layout="wide") # 넓은 레이아웃 사용
st.title("📄 문서 자동화 시연 프로토타입")
st.markdown("---")

# --- 프로젝트 기본 정보 입력 섹션 ---
st.header("1. 프로젝트 기본 정보 입력")
col1, col2 = st.columns(2) # 두 개의 컬럼으로 레이아웃 분할

with col1:
    project_title = st.text_input("프로젝트 제목", "새로운 AI 헬스케어 시스템", help="생성될 문서의 주요 제목입니다.")
    author = st.text_input("작성자", "팀 헬띠", help="문서 작성자의 이름 또는 팀명입니다.")

with col2:
    current_date = datetime.now().strftime("%Y-%m-%d")
    doc_date = st.text_input("문서 작성일", current_date, help="문서에 표시될 날짜입니다.")
    summary = st.text_area("프로젝트 요약", "사용자 맞춤형 운동 추천 및 건강 분석을 위한 AI 기반 서비스 프로토타입입니다.", help="프로젝트의 간략한 설명을 입력하세요.")

st.markdown("---")

# --- 요구사항 입력 및 관리 섹션 ---
st.header("2. 요구사항 입력 및 관리")

# 세션 상태에 'requirements' DataFrame이 없으면 초기화
if 'requirements' not in st.session_state:
    st.session_state.requirements = pd.DataFrame(columns=['요구사항ID', '분류', '요구사항 명', '요구사항 설명'])

req_name = st.text_input("요구사항 명칭", placeholder="예: 운동 추천 기능 구현", key="req_name_input", help="요구사항의 핵심 기능을 나타내는 제목입니다.")
req_description = st.text_area("요구사항 상세 설명", placeholder="예: 사용자의 건강 데이터를 분석하여 맞춤형 운동 종목을 추천합니다.", key="req_desc_input", help="요구사항에 대한 구체적인 설명입니다.")
req_category = st.radio("요구사항 분류", ["기능", "비기능"], horizontal=True, key="req_category_input", help="요구사항이 기능(Functional)인지 비기능(Non-Functional)인지 선택하세요.")

if st.button("➕ 요구사항 추가", help="위 입력된 요구사항을 목록에 추가합니다."):
    if req_name and req_description:
        new_id = generate_requirement_id(req_category) # 자동 ID 생성
        new_requirement = pd.DataFrame([{
            '요구사항ID': new_id,
            '분류': req_category,
            '요구사항 명': req_name,
            '요구사항 설명': req_description
        }])
        # 기존 DataFrame에 새로운 요구사항 추가
        st.session_state.requirements = pd.concat([st.session_state.requirements, new_requirement], ignore_index=True)
        st.success(f"요구사항 '{req_name}' (ID: {new_id})이(가) 성공적으로 추가되었습니다.")
        # 입력 필드 초기화 (사용자 경험 개선)
        st.session_state["req_name_input"] = ""
        st.session_state["req_desc_input"] = ""
    else:
        st.warning("🚨 요구사항 명칭과 상세 설명을 모두 입력해주세요.")

st.subheader("현재 등록된 요구사항 목록")
st.dataframe(st.session_state.requirements, use_container_width=True)

# 요구사항 목록 초기화 버튼 (시연 용이성)
if st.button("🔄 요구사항 목록 초기화", help="등록된 모든 요구사항을 삭제하고 ID 카운터를 초기화합니다."):
    st.session_state.requirements = pd.DataFrame(columns=['요구사항ID', '분류', '요구사항 명', '요구사항 설명'])
    st.session_state.next_id_f = 1
    st.session_state.next_id_nf = 1
    st.info("요구사항 목록이 초기화되었습니다.")

st.markdown("---")

# --- 문서 생성 및 다운로드 섹션 ---
st.header("3. 문서 생성 및 다운로드")

if st.button("🚀 문서 생성 및 다운로드", help="입력된 정보와 요구사항을 기반으로 Word 문서를 생성합니다."):
    project_info = {
        "project_title": project_title,
        "author": author,
        "date": doc_date,
        "summary": summary
    }
    
    doc_buffer = create_document(project_info, st.session_state.requirements)
    
    # 생성된 문서를 다운로드 버튼으로 제공
    st.download_button(
        label="📥 설계서.docx 다운로드",
        data=doc_buffer,
        file_name="생성된_설계서.docx", # 다운로드될 파일명
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="클릭하여 생성된 Word 문서를 다운로드합니다."
    )
    st.success("🎉 문서가 성공적으로 생성되었습니다! 위 다운로드 버튼을 클릭하세요.")
else:
    st.info("⬆️ 프로젝트 정보와 요구사항을 입력한 후 '문서 생성 및 다운로드' 버튼을 클릭하세요.")