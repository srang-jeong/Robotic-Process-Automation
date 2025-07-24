import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import pandas as pd
from datetime import datetime, timedelta
import zipfile
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.font_manager as fm
import numpy as np
from matplotlib.ticker import AutoLocator

# --- 폰트 설정 (한글 깨짐 방지) ---
try:
    font_path = fm.findfont(fm.FontProperties(family='NanumGothic'))
    plt.rcParams['font.family'] = 'NanumGothic'
except:
    plt.rcParams['font.family'] = 'Malgun Gothic'
    if 'Malgun Gothic' not in fm.get_font_names():
        plt.rcParams['font.family'] = 'AppleGothic'
        if 'AppleGothic' not in fm.get_font_names():
             st.warning("경고: 그래프의 한글 폰트를 찾을 수 없습니다. 그래프의 한글이 깨져 보일 수 있습니다. 'NanumGothic' 또는 'Malgun Gothic' 폰트 설치를 권장합니다.")
             plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

# --- Word 문서 한글 폰트 설정 함수 ---
def set_korean_font(document):
    # 기본 스타일 폰트 설정
    document.styles['Normal'].font.name = '맑은 고딕' # 기본 한글 폰트
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕') # 동아시아 폰트 설정
    
    # Heading 1, 2, 3 등에도 적용 (필요시)
    for i in range(10): # Heading 0 to Heading 9
        if f'Heading {i}' in document.styles:
            document.styles[f'Heading {i}'].font.name = '맑은 고딕'
            document.styles[f'Heading {i}']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# --- 1. ID 생성 로직 (기존과 동일) ---
def generate_requirement_id(category):
    # 이 함수 외부에서 session_state 초기화가 이루어지므로, 여기서는 바로 사용
    if category == "기능":
        current_id = st.session_state.next_id_f
        st.session_state.next_id_f += 1
        return f"RQ-F-{current_id:03d}"
    elif category == "비기능":
        current_id = st.session_state.next_id_nf
        st.session_state.next_id_nf += 1
        return f"RQ-NF-{current_id:03d}"
    else:
        return "N/A"

# --- 2. 문서 생성 로직 (Word 문서) ---

def create_project_proposal_document(project_info, doc_history_df):
    document = Document()
    set_korean_font(document) # 한글 폰트 설정 적용

    document.add_heading(f'{project_info["project_title"]} 프로젝트 기획서', level=0)
    document.add_paragraph(f'작성자: {project_info["author"]}')
    document.add_paragraph(f'작성일: {project_info["date"]}')
    document.add_page_break()

    # --- 문서 이력 섹션 추가 ---
    document.add_heading('0. 문서 이력', level=1)
    if not doc_history_df.empty:
        table = document.add_table(rows=1, cols=len(doc_history_df.columns))
        table.style = 'Table Grid'
        
        # 헤더 셀 설정
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(doc_history_df.columns):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True

        # 데이터 행 추가
        for index, row in doc_history_df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(doc_history_df.columns):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 문서 변경 이력이 없습니다.")
    document.add_page_break()
    # --- 문서 이력 섹션 끝 ---


    document.add_heading('1. 개요', level=1)
    document.add_paragraph(f'본 문서는 "{project_info["project_title"]}" 프로젝트의 기획을 위한 문서입니다. 이 프로젝트는 {project_info["summary"]}와 같은 목표를 가지고 있습니다.')
    
    document.add_heading('2. 개발 목표', level=1)
    document.add_paragraph('프로젝트의 주요 개발 목표와 기대 효과를 상세하게 기술합니다.')
    document.add_paragraph('  - 목표 1: (여기에 목표 내용 추가)')
    document.add_paragraph('  - 목표 2: (여기에 목표 내용 추가)')

    document.add_heading('3. 범위', level=1)
    document.add_paragraph('프로젝트의 포함 범위 및 제외되는 범위를 명확히 정의합니다.')
    document.add_paragraph('  - 포함 범위: (내용 추가)')
    document.add_paragraph('  - 제외 범위: (내용 추가)')

    document.add_heading('4. 추진 일정 (간략)', level=1)
    document.add_paragraph('프로젝트의 주요 단계별 일정을 간략하게 기술합니다.')
    document.add_paragraph('  - 1단계: 기획 및 분석')
    document.add_paragraph('  - 2단계: 설계 및 개발')
    document.add_paragraph('  - 3단계: 테스트 및 배포')

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_project_result_report_document(project_info, doc_history_df):
    document = Document()
    set_korean_font(document) # 한글 폰트 설정 적용

    document.add_heading(f'{project_info["project_title"]} 프로젝트 결과 보고서', level=0)
    document.add_paragraph(f'작성자: {project_info["author"]}')
    document.add_paragraph(f'작성일: {project_info["date"]}')
    document.add_page_break()

    # --- 문서 이력 섹션 추가 ---
    document.add_heading('0. 문서 이력', level=1)
    if not doc_history_df.empty:
        table = document.add_table(rows=1, cols=len(doc_history_df.columns))
        table.style = 'Table Grid'
        
        # 헤더 셀 설정
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(doc_history_df.columns):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True

        # 데이터 행 추가
        for index, row in doc_history_df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(doc_history_df.columns):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 문서 변경 이력이 없습니다.")
    document.add_page_break()
    # --- 문서 이력 섹션 끝 ---

    document.add_heading('1. 프로젝트 성과 요약', level=1)
    document.add_paragraph(f'"{project_info["project_title"]}" 프로젝트의 주요 성과와 달성률을 요약합니다.')
    document.add_paragraph('  - 주요 달성 목표: (내용 추가)')
    document.add_paragraph('  - 정량적 성과: (내용 추가)')

    document.add_heading('2. 구현 기능 및 결과', level=1)
    document.add_paragraph('각 기능별 구현 내용 및 테스트 결과를 기술합니다.')
    document.add_paragraph('  - 기능 A: (구현 내용 및 결과)')
    document.add_paragraph('  - 기능 B: (구현 내용 및 결과)')

    document.add_heading('3. 문제점 및 개선 방안', level=1)
    document.add_paragraph('프로젝트 진행 중 발생했던 문제점과 이에 대한 해결 또는 개선 방안을 제시합니다.')
    document.add_paragraph('  - 발생 문제: (내용 추가)')
    document.add_paragraph('  - 개선 방안: (내용 추가)')

    document.add_heading('4. 향후 계획', level=1)
    document.add_paragraph('프로젝트의 다음 단계 또는 추가 개발 계획을 기술합니다.')
    document.add_paragraph('  - 1차 목표: (내용 추가)')
    document.add_paragraph('  - 장기 목표: (내용 추가)')

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_requirements_document(project_info, requirements_df, doc_history_df):
    document = Document()
    set_korean_font(document) # 한글 폰트 설정 적용

    document.add_heading(f'{project_info["project_title"]} 요구사항 명세서', level=0)
    document.add_paragraph(f'작성일: {project_info["date"]}')
    document.add_page_break()

    # --- 문서 이력 섹션 추가 ---
    document.add_heading('0. 문서 이력', level=1)
    if not doc_history_df.empty:
        table = document.add_table(rows=1, cols=len(doc_history_df.columns))
        table.style = 'Table Grid'
        
        # 헤더 셀 설정
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(doc_history_df.columns):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True

        # 데이터 행 추가
        for index, row in doc_history_df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(doc_history_df.columns):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 문서 변경 이력이 없습니다.")
    document.add_page_break()
    # --- 문서 이력 섹션 끝 ---

    document.add_heading('1. 요구사항 목록', level=1)

    if not requirements_df.empty:
        # 요구사항 명세서에는 중복 제거된 요구사항만 표시 (테스트 케이스는 추적표에서)
        # '요구사항ID', '요구사항 명', '요구사항 설명' 기준으로 중복 제거
        unique_reqs_df = requirements_df.drop_duplicates(subset=['요구사항ID', '요구사항 명', '요구사항 설명']).copy() # .copy() 추가
        
        display_cols = ['요구사항ID', '분류', '요구사항 명', '요구사항 설명']
        table = document.add_table(rows=1, cols=len(display_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(display_cols):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True

        for index, row in unique_reqs_df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(display_cols):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 요구사항이 없습니다.")
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_requirements_traceability_matrix_document(project_info, requirements_df, doc_history_df):
    document = Document()
    set_korean_font(document) # 한글 폰트 설정 적용

    document.add_heading(f'{project_info["project_title"]} 요구사항 추적표', level=0)
    document.add_paragraph(f'작성일: {project_info["date"]}')
    document.add_page_break()

    # --- 문서 이력 섹션 추가 ---
    document.add_heading('0. 문서 이력', level=1)
    if not doc_history_df.empty:
        table = document.add_table(rows=1, cols=len(doc_history_df.columns))
        table.style = 'Table Grid'
        
        # 헤더 셀 설정
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(doc_history_df.columns):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True

        # 데이터 행 추가
        for index, row in doc_history_df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(doc_history_df.columns):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 문서 변경 이력이 없습니다.")
    document.add_page_break()
    # --- 문서 이력 섹션 끝 ---

    document.add_heading('1. 요구사항 추적표', level=1)
    document.add_paragraph('이 표는 등록된 요구사항과 관련된 테스트 항목을 추적합니다.')

    if not requirements_df.empty:
        trace_cols = [
            '요구사항ID', '요구사항 명', '요구사항 설명',
            '테스트케이스ID', '테스트케이스 명', '테스트케이스 설명', '테스트 방법'
        ]
        
        table = document.add_table(rows=1, cols=len(trace_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(trace_cols):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs: # 헤더 볼드체
                for run in paragraph.runs:
                    run.bold = True
        
        for index, row in requirements_df.iterrows(): # 모든 행을 순회 (1:N 관계 반영)
            row_cells = table.add_row().cells
            for i, col_name in enumerate(trace_cols):
                row_cells[i].text = str(row.get(col_name, ''))
    else:
        document.add_paragraph("등록된 요구사항이 없어 추적표를 생성할 수 없습니다.")
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. 일정 관련 문서 생성 함수 (간트 차트 PDF, 일정표 XLS) ---

def create_gantt_chart_pdf(project_info, schedule_df, show_dependencies=False):
    fig, ax = plt.subplots(figsize=(12, len(schedule_df) * 0.7 + 2))

    if not schedule_df.empty:
        schedule_df['Start Date'] = pd.to_datetime(schedule_df['Start Date'], errors='coerce')
        schedule_df['End Date'] = pd.to_datetime(schedule_df['End Date'], errors='coerce')
        
        # schedule_df = schedule_df.dropna(subset=['Start Date', 'End Date'], inplace=True) # 변경된 부분
        schedule_df = schedule_df.dropna(subset=['Start Date', 'End Date']).copy() # .copy() 추가

        if schedule_df.empty:
            ax.text(0.5, 0.5, '유효한 일정 데이터가 없어 간트 차트를 그릴 수 없습니다.', 
                    horizontalalignment='center', verticalalignment='center', 
                    transform=ax.transAxes, fontsize=12, color='gray')
            ax.set_axis_off()
        else:
            schedule_df['Duration'] = (schedule_df['End Date'] - schedule_df['Start Date']).dt.days

            tasks = []
            for index, row in schedule_df.iterrows():
                task_label = row['Task Name']
                if 'Coordinator' in schedule_df.columns and pd.notna(row['Coordinator']) and row['Coordinator'].strip():
                    task_label += f" ({row['Coordinator']})" # 담당자 추가
                if show_dependencies and 'Predecessors' in schedule_df.columns and pd.notna(row['Predecessors']) and row['Predecessors'].strip():
                    task_label += f"\n(선행: {row['Predecessors']})"
                tasks.append(task_label)

            y_pos = np.arange(len(tasks))

            # 색상 및 스타일 개선
            colors = plt.cm.viridis(y_pos / len(y_pos)) # 작업별 다른 색상
            ax.barh(y_pos, schedule_df['Duration'], left=mdates.date2num(schedule_df['Start Date']), 
                    height=0.6, align='center', color=colors, edgecolor='black', linewidth=0.7)

            ax.set_yticks(y_pos)
            ax.set_yticklabels(tasks, fontsize=10)
            ax.invert_yaxis()

            ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
            ax.xaxis.set_major_locator(AutoLocator())
            fig.autofmt_xdate()

            ax.set_title(f'{project_info["project_title"]} 프로젝트 간트 차트', fontsize=16, pad=15)
            ax.set_xlabel('날짜', fontsize=12)
            ax.grid(True, linestyle='--', alpha=0.6)
            ax.tick_params(axis='x', rotation=45) # x축 텍스트 회전
            
            # 그래프 외곽선 제거
            for spine in ax.spines.values():
                spine.set_visible(False)

    else:
        ax.text(0.5, 0.5, '등록된 일정 데이터가 없습니다.', 
                horizontalalignment='center', verticalalignment='center', 
                transform=ax.transAxes, fontsize=12, color='gray')
        ax.set_axis_off()

    plt.tight_layout()
    
    buffer = io.BytesIO()
    plt.savefig(buffer, format='pdf', bbox_inches='tight')
    plt.close(fig)
    buffer.seek(0)
    return buffer

def create_schedule_excel(project_info, schedule_df):
    buffer = io.BytesIO()
    if not schedule_df.empty:
        # 엑셀 파일에 스타일 적용 (예: 헤더 볼드체)
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            schedule_df.to_excel(writer, index=False, sheet_name='프로젝트 일정')
            workbook = writer.book
            worksheet = writer.sheets['프로젝트 일정']
            
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#D7E4BC', # 연두색 배경
                'border': 1
            })

            for col_num, value in enumerate(schedule_df.columns.values):
                worksheet.write(0, col_num, value, header_format)
            
            # 컬럼 너비 자동 조정 (간단한 방법)
            for i, col in enumerate(schedule_df.columns):
                max_len = max(schedule_df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, max_len)
    else:
        # 'Coordinator' 컬럼도 포함하여 초기화
        empty_df = pd.DataFrame(columns=['Task Name', 'Start Date', 'End Date', 'Predecessors', 'Coordinator'])
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            empty_df.to_excel(writer, index=False, sheet_name='프로젝트 일정')
            workbook = writer.book
            worksheet = writer.sheets['프로젝트 일정']
            header_format = workbook.add_format({'bold': True, 'fg_color': '#D7E4BC', 'border': 1})
            for col_num, value in enumerate(empty_df.columns.values):
                worksheet.write(0, col_num, value, header_format)
            for i, col in enumerate(empty_df.columns):
                max_len = len(col) + 2
                worksheet.set_column(i, i, max_len)
    buffer.seek(0)
    return buffer

# --- 4. Streamlit UI 구성 ---
st.set_page_config(layout="wide")
st.title("📄 문서 자동화 시연 프로토타입")
st.markdown("---")

# --- 세션 상태 변수 초기화 (가장 상단에 위치) ---
if 'next_id_f' not in st.session_state:
    st.session_state.next_id_f = 1
if 'next_id_nf' not in st.session_state:
    st.session_state.next_id_nf = 1
if 'requirements' not in st.session_state:
    st.session_state.requirements = pd.DataFrame(columns=[
        '요구사항ID', '분류', '요구사항 명', '요구사항 설명',
        '테스트케이스ID', '테스트케이스 명', '테스트케이스 설명', '테스트 방법'
    ])
if 'schedule_data' not in st.session_state:
    st.session_state.schedule_data = pd.DataFrame(columns=['Task Name', 'Start Date', 'End Date', 'Predecessors', 'Coordinator'])
if 'document_history' not in st.session_state:
    st.session_state.document_history = pd.DataFrame(columns=['버전', '내용', '작성자', '일자'])


# --- 프로젝트 기본 정보 입력 섹션 ---
st.header("1. 프로젝트 기본 정보 입력")
col1, col2 = st.columns(2)

with col1:
    project_title = st.text_input("프로젝트 제목", "새로운 AI 헬스케어 시스템", help="생성될 문서의 주요 제목입니다.")
    author = st.text_input("작성자", "김프로토", help="문서 작성자의 이름 또는 팀명입니다.")

with col2:
    current_date = datetime.now().strftime("%Y-%m-%d")
    doc_date = st.text_input("문서 작성일", current_date, help="문서에 표시될 날짜입니다.")
    summary = st.text_area("프로젝트 요약", "사용자 맞춤형 운동 추천 및 건강 분석을 위한 AI 기반 서비스 프로토타입입니다.", help="프로젝트의 간략한 설명을 입력하세요.")

st.markdown("---")

# --- 요구사항 입력 및 관리 섹션 ---
st.header("2. 요구사항 입력 및 관리")

# 요구사항 추가 모드 선택
# 세션 상태 초기화
if 'next_id_f' not in st.session_state:
    st.session_state.next_id_f = 1
if 'next_id_nf' not in st.session_state:
    st.session_state.next_id_nf = 1
if 'requirements' not in st.session_state:
    st.session_state.requirements = pd.DataFrame(columns=[
        '요구사항ID', '분류', '요구사항 명', '요구사항 설명',
        '테스트케이스ID', '테스트케이스 명', '테스트케이스 설명', '테스트 방법'
    ])
if 'document_history' not in st.session_state:
    st.session_state.document_history = pd.DataFrame(columns=['버전', '내용', '작성자', '일자'])

# 요구사항 ID 생성 함수
def generate_requirement_id(category):
    if category == "기능":
        current_id = st.session_state.next_id_f
        st.session_state.next_id_f += 1
        return f"RQ-F-{current_id:03d}"
    elif category == "비기능":
        current_id = st.session_state.next_id_nf
        st.session_state.next_id_nf += 1
        return f"RQ-NF-{current_id:03d}"
    else:
        return "N/A"

st.header("2. 요구사항 입력 및 관리")

add_mode = st.radio(
    "요구사항 추가 방식 선택",
    ["새로운 요구사항 추가", "기존 요구사항에 테스트케이스 추가"],
    horizontal=True
)

if add_mode == "새로운 요구사항 추가":
    st.subheader("새로운 요구사항 등록")
    with st.form("new_req_form", clear_on_submit=True):
        req_name = st.text_input("요구사항 명칭", placeholder="예: 운동 추천 기능 구현")
        req_description = st.text_area("요구사항 상세 설명", placeholder="예: 사용자 건강 데이터 분석, 맞춤형 운동 추천")
        req_category = st.radio("요구사항 분류", ["기능", "비기능"], horizontal=True)
        st.subheader("테스트 관련 정보 (선택 사항 - 첫 번째 테스트 케이스)")
        tc_id = st.text_input("테스트케이스 ID", placeholder="예: TC-RQF-001-01")
        tc_name = st.text_input("테스트케이스 명칭", placeholder="예: 운동 추천 기능 단위 테스트")
        tc_description = st.text_area("테스트케이스 상세 설명", placeholder="예: 입력 → 추천 목록 확인")
        tc_method = st.selectbox("테스트 방법", ["선택 안함", "단위 테스트", "통합 테스트", "시스템 테스트", "인수 테스트"])
        submitted = st.form_submit_button("➕ 새로운 요구사항 추가")

    if submitted:
        if req_name and req_description:
            new_id = generate_requirement_id(req_category)
            new_row = pd.DataFrame([{
                '요구사항ID': new_id,
                '분류': req_category,
                '요구사항 명': req_name,
                '요구사항 설명': req_description,
                '테스트케이스ID': tc_id if tc_id else None,
                '테스트케이스 명': tc_name if tc_name else None,
                '테스트케이스 설명': tc_description if tc_description else None,
                '테스트 방법': tc_method if tc_method != "선택 안함" else None
            }])
            st.session_state.requirements = pd.concat([st.session_state.requirements, new_row], ignore_index=True)
            st.success(f"요구사항 '{req_name}' (ID: {new_id})이(가) 성공적으로 추가되었습니다.")
        else:
            st.warning("🚨 요구사항 명칭과 상세 설명을 모두 입력해주세요.")
elif add_mode == "기존 요구사항에 테스트케이스 추가":
    st.subheader("기존 요구사항에 테스트케이스 추가")
    if not st.session_state.requirements.empty:
        unique_req_options = st.session_state.requirements[['요구사항ID', '요구사항 명']].drop_duplicates().copy()
        unique_req_options['Display'] = unique_req_options['요구사항 명'] + " (" + unique_req_options['요구사항ID'] + ")"
        selected_req_display = st.selectbox(
            "테스트케이스를 추가할 요구사항 선택",
            unique_req_options['Display']
        )
        selected_req_id = unique_req_options[unique_req_options['Display'] == selected_req_display]['요구사항ID'].iloc[0]
        selected_req_info = st.session_state.requirements[st.session_state.requirements['요구사항ID'] == selected_req_id].iloc[0]
        st.markdown(f"**선택된 요구사항:** {selected_req_info['요구사항 명']} ({selected_req_info['요구사항ID']})")

        with st.form("add_tc_form", clear_on_submit=True):
            tc_id = st.text_input("테스트케이스 ID", placeholder=f"예: TC-{selected_req_id}-02")
            tc_name = st.text_input("테스트케이스 명칭", placeholder="예: 운동 추천 기능 통합 테스트")
            tc_description = st.text_area("테스트케이스 상세 설명", placeholder="테스트케이스 설명 입력")
            tc_method = st.selectbox("테스트 방법", ["선택 안함", "단위 테스트", "통합 테스트", "시스템 테스트", "인수 테스트"])
            tc_submitted = st.form_submit_button("➕ 테스트케이스 추가")
        if tc_submitted:
            if tc_id and tc_name and tc_description and tc_method != "선택 안함":
                new_row = pd.DataFrame([{
                    '요구사항ID': selected_req_info['요구사항ID'],
                    '분류': selected_req_info['분류'],
                    '요구사항 명': selected_req_info['요구사항 명'],
                    '요구사항 설명': selected_req_info['요구사항 설명'],
                    '테스트케이스ID': tc_id,
                    '테스트케이스 명': tc_name,
                    '테스트케이스 설명': tc_description,
                    '테스트 방법': tc_method
                }])
                st.session_state.requirements = pd.concat([st.session_state.requirements, new_row], ignore_index=True)
                st.success(f"테스트케이스 '{tc_name}'이(가) 추가되었습니다.")
            else:
                st.warning("🚨 모든 테스트케이스 정보를 입력해주세요.")
    else:
        st.info("💡 등록된 요구사항이 없습니다.")

st.subheader("현재 등록된 요구사항 및 테스트케이스 목록")
st.dataframe(st.session_state.requirements, use_container_width=True)

if st.button("🔄 요구사항 목록 초기화"):
    st.session_state.requirements = pd.DataFrame(columns=[
        '요구사항ID', '분류', '요구사항 명', '요구사항 설명',
        '테스트케이스ID', '테스트케이스 명', '테스트케이스 설명', '테스트 방법'
    ])
    st.session_state.next_id_f = 1
    st.session_state.next_id_nf = 1
    st.info("요구사항 및 테스트케이스 목록이 초기화되었습니다.")

st.markdown("---")

# --- 일정 관리 섹션 ---
st.header("3. 프로젝트 일정 관리")

st.markdown("##### 3.1. 일정 CSV/Excel 파일 업로드")
st.info("💡 일정 파일은 'Task Name', 'Start Date', 'End Date', 그리고 선택적으로 'Predecessors' 및 **'Coordinator'** 컬럼을 포함해야 합니다. 날짜 형식은 YYYY-MM-DD, Predecessors는 쉼표로 구분된 선행 작업 이름을 입력해주세요.")
uploaded_schedule_file = st.file_uploader("일정 데이터를 포함하는 CSV 또는 Excel 파일을 업로드하세요.", type=["csv", "xlsx"])


if uploaded_schedule_file is not None:
    try:
        if uploaded_schedule_file.name.endswith('.csv'):
            temp_df = pd.read_csv(uploaded_schedule_file)
        else: # .xlsx
            temp_df = pd.read_excel(uploaded_schedule_file)
        
        required_cols = ['Task Name', 'Start Date', 'End Date']
        if all(col in temp_df.columns for col in required_cols):
            # 'Predecessors'와 'Coordinator' 컬럼이 없으면 추가 (None으로 채워짐)
            if 'Predecessors' not in temp_df.columns:
                temp_df['Predecessors'] = None
            if 'Coordinator' not in temp_df.columns: # 새로 추가된 컬럼
                temp_df['Coordinator'] = None
            
            st.session_state.schedule_data = temp_df[required_cols + ['Predecessors', 'Coordinator']].copy()
            st.success("✅ 일정 데이터가 성공적으로 로드되었습니다.")
        else:
            st.error(f"⚠️ 업로드된 파일에 필수 컬럼({', '.join(required_cols)})이 모두 포함되어 있지 않습니다.")
    except Exception as e:
        st.error(f"파일 로드 중 오류가 발생했습니다: {e}")

st.subheader("현재 등록된 일정")
st.dataframe(st.session_state.schedule_data, use_container_width=True)

if st.button("🔄 일정 데이터 초기화", help="등록된 일정 데이터를 모두 삭제합니다."):
    # 'Coordinator' 컬럼도 포함하여 초기화
    st.session_state.schedule_data = pd.DataFrame(columns=['Task Name', 'Start Date', 'End Date', 'Predecessors', 'Coordinator'])
    st.info("일정 데이터가 초기화되었습니다.")

st.markdown("---")

# --- 문서 변경 이력 관리 섹션 ---
st.header("4. 문서 변경 이력 관리")
st.subheader("새로운 변경 이력 추가")

col_hist1, col_hist2 = st.columns(2)
with st.form("doc_hist_form", clear_on_submit=True):
    with col_hist1:
        hist_version = st.text_input("버전", "1.0", key="hist_version_input", help="문서 버전 (예: 1.0, 1.1)")
        hist_author = st.text_input("작성자", author, key="hist_author_input", help="변경 내용을 작성한 사람")
    with col_hist2:
        hist_date = st.text_input("일자", datetime.now().strftime("%Y-%m-%d"), key="hist_date_input", help="변경이 발생한 날짜")
        hist_content = st.text_area("내용", "초기 문서 작성", key="hist_content_input", help="변경 내용 요약")
    submitted = st.form_submit_button("➕ 변경 이력 추가")

if submitted:
    if hist_version and hist_content and hist_author and hist_date:
        new_row = pd.DataFrame([{
            '버전': hist_version,
            '내용': hist_content,
            '작성자': hist_author,
            '일자': hist_date
        }])
        st.session_state.document_history = pd.concat([st.session_state.document_history, new_row], ignore_index=True)
        st.success("변경 이력이 추가되었습니다.")
        # 버전 자동증가는 메시지로 안내만
        try:
            next_version = str(round(float(hist_version) + 0.1, 1))
        except:
            next_version = "1.1"
        st.info(f"다음 버전(수동 입력): {next_version}")
    else:
        st.warning("🚨 모든 변경 이력 필드를 채워주세요.")


st.subheader("현재 등록된 문서 변경 이력")
if not st.session_state.document_history.empty:
    st.dataframe(st.session_state.document_history, use_container_width=True)
else:
    st.info("등록된 문서 변경 이력이 없습니다.")

if st.button("🔄 변경 이력 목록 초기화", help="등록된 모든 변경 이력 데이터를 삭제합니다."):
    st.session_state.document_history = pd.DataFrame(columns=['버전', '내용', '작성자', '일자'])
    st.info("문서 변경 이력 목록이 초기화되었습니다.")

st.markdown("---")


# --- 최종 문서 생성 및 다운로드 섹션 ---
st.header("5. 모든 문서 생성 및 다운로드") # 섹션 번호 변경됨

# project_info 딕셔너리를 여기서 정의하여, 위의 UI 입력 값들을 최신으로 반영하도록 합니다.
project_info = {
    "project_title": project_title,
    "author": author,
    "date": doc_date,
    "summary": summary
}

# --- 모든 문서 ZIP으로 다운로드 ---
st.subheader("5.1. 모든 문서를 ZIP 파일로 다운로드")
show_gantt_dependencies = st.checkbox("간트 차트에 선행 의존 정보 표시", value=False, help="체크하면 간트 차트의 작업 이름 옆에 선행 작업 정보가 표시됩니다.")

# ZIP 파일 생성 로직은 버튼이 눌렸을 때만 실행되도록 콜백 함수 또는 if 문으로 유지
if st.button("📦 모든 문서 압축 다운로드", help="프로젝트 기획서, 결과 보고서, 요구사항 명세서, 요구사항 추적표, 간트 차트(PDF), 일정표(XLS)를 생성하여 하나의 ZIP 파일로 다운로드합니다."):
    
    # 여기서 모든 문서 버퍼를 생성
    proposal_doc_buffer = create_project_proposal_document(project_info, st.session_state.document_history.copy())
    result_report_doc_buffer = create_project_result_report_document(project_info, st.session_state.document_history.copy())
    reqs_spec_doc_buffer = create_requirements_document(project_info, st.session_state.requirements.copy(), st.session_state.document_history.copy())
    reqs_trace_doc_buffer = create_requirements_traceability_matrix_document(project_info, st.session_state.requirements.copy(), st.session_state.document_history.copy())

    gantt_chart_pdf_buffer = create_gantt_chart_pdf(project_info, st.session_state.schedule_data.copy(), show_dependencies=show_gantt_dependencies)
    schedule_excel_buffer = create_schedule_excel(project_info, st.session_state.schedule_data.copy())

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{project_info['project_title']}_프로젝트_기획서.docx", proposal_doc_buffer.getvalue())
        zf.writestr(f"{project_info['project_title']}_프로젝트_결과보고서.docx", result_report_doc_buffer.getvalue())
        zf.writestr(f"{project_info['project_title']}_요구사항_명세서.docx", reqs_spec_doc_buffer.getvalue())
        zf.writestr(f"{project_info['project_title']}_요구사항_추적표.docx", reqs_trace_doc_buffer.getvalue())
        
        zf.writestr(f"{project_info['project_title']}_간트차트.pdf", gantt_chart_pdf_buffer.getvalue())
        zf.writestr(f"{project_info['project_title']}_일정표.xlsx", schedule_excel_buffer.getvalue())

    zip_buffer.seek(0)

    st.download_button(
        label="📥 모든 문서 ZIP 파일 다운로드",
        data=zip_buffer,
        file_name=f"{project_title}_문서_패키지.zip",
        mime="application/zip",
        help="생성된 모든 문서(Word, PDF, Excel)가 포함된 ZIP 파일을 다운로드합니다."
    )
    st.success("🎉 모든 문서가 성공적으로 생성되어 ZIP 파일로 압축되었습니다! 위 버튼을 클릭하여 다운로드하세요.")
else:
    st.info("⬆️ 프로젝트 정보, 요구사항, 일정을 입력한 후 '모든 문서 압축 다운로드' 또는 개별 다운로드 버튼을 클릭하세요.")


# --- 개별 문서 다운로드 ---
st.subheader("5.2. 각 문서를 개별적으로 다운로드") # 섹션 번호 변경됨

col_buttons = st.columns(3) # 3개의 컬럼으로 버튼 배치

# 프로젝트 기획서
with col_buttons[0]:
    # Streamlit의 download_button은 if 문으로 묶을 필요가 없습니다.
    # 파일 버퍼를 생성하는 함수를 직접 data 인수에 전달하거나, 미리 버퍼를 생성합니다.
    st.download_button(
        label="📝 기획서 다운로드",
        data=create_project_proposal_document(project_info, st.session_state.document_history.copy()),
        file_name=f"{project_info['project_title']}_프로젝트_기획서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_proposal_button", # 키 변경 (이전 키는 if 문 안에 있었으므로)
        help="프로젝트 기획서 문서를 다운로드합니다."
    )

# 프로젝트 결과 보고서
with col_buttons[1]:
    st.download_button(
        label="📊 결과 보고서 다운로드",
        data=create_project_result_report_document(project_info, st.session_state.document_history.copy()),
        file_name=f"{project_info['project_title']}_프로젝트_결과보고서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_result_report_button",
        help="프로젝트 결과 보고서 문서를 다운로드합니다."
    )

# 요구사항 명세서
with col_buttons[2]:
    st.download_button(
        label="📑 요구사항 명세서 다운로드",
        data=create_requirements_document(project_info, st.session_state.requirements.copy(), st.session_state.document_history.copy()),
        file_name=f"{project_info['project_title']}_요구사항_명세서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_reqs_spec_button",
        help="요구사항 명세서 문서를 다운로드합니다."
    )

# 새로운 줄에 버튼 배치
col_buttons_2 = st.columns(3)

# 요구사항 추적표
with col_buttons_2[0]:
    st.download_button(
        label="🔍 요구사항 추적표 다운로드",
        data=create_requirements_traceability_matrix_document(project_info, st.session_state.requirements.copy(), st.session_state.document_history.copy()),
        file_name=f"{project_info['project_title']}_요구사항_추적표.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_reqs_trace_button",
        help="요구사항 추적표 문서를 다운로드합니다."
    )

# 간트 차트 PDF
with col_buttons_2[1]:
    st.download_button(
        label="🗓️ 간트 차트 PDF 다운로드",
        data=create_gantt_chart_pdf(project_info, st.session_state.schedule_data.copy(), show_dependencies=show_gantt_dependencies),
        file_name=f"{project_info['project_title']}_간트차트.pdf",
        mime="application/pdf",
        key="download_gantt_pdf_button",
        help="간트 차트를 PDF 형식으로 다운로드합니다."
    )

# 일정표 Excel
with col_buttons_2[2]:
    st.download_button(
        label="🗓️ 일정표 Excel 다운로드",
        data=create_schedule_excel(project_info, st.session_state.schedule_data.copy()),
        file_name=f"{project_info['project_title']}_일정표.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="download_schedule_excel_button",
        help="프로젝트 일정표를 Excel 형식으로 다운로드합니다."
    )

st.markdown("---")

# --- 6. 데이터 저장 및 불러오기 섹션 (번호 변경됨) ---
st.header("6. 데이터 아카이빙 및 불러오기")

st.markdown("##### 6.1. 현재 데이터 저장")
st.info("💡 현재 입력된 요구사항/테스트케이스 목록, 일정 데이터 및 문서 이력 데이터를 ZIP 파일로 저장합니다. 나중에 이 파일을 업로드하여 데이터를 복원할 수 있습니다.")

if st.button("💾 현재 데이터 ZIP으로 저장", help="현재 입력된 요구사항/테스트케이스 목록, 일정 데이터 및 문서 이력 데이터를 압축 파일로 저장합니다."):
    if not st.session_state.requirements.empty or not st.session_state.schedule_data.empty or not st.session_state.document_history.empty:
        data_zip_buffer = io.BytesIO()
        with zipfile.ZipFile(data_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 요구사항 데이터 CSV로 저장
            req_csv_buffer = io.StringIO()
            st.session_state.requirements.to_csv(req_csv_buffer, index=False, encoding='utf-8-sig')
            zf.writestr("requirements_data.csv", req_csv_buffer.getvalue())
            
            # 일정 데이터 CSV로 저장 ('Coordinator' 컬럼 포함)
            schedule_csv_buffer = io.StringIO()
            st.session_state.schedule_data.to_csv(schedule_csv_buffer, index=False, encoding='utf-8-sig')
            zf.writestr("schedule_data.csv", schedule_csv_buffer.getvalue())

            # 문서 이력 데이터 CSV로 저장
            history_csv_buffer = io.StringIO()
            st.session_state.document_history.to_csv(history_csv_buffer, index=False, encoding='utf-8-sig')
            zf.writestr("document_history_data.csv", history_csv_buffer.getvalue())
            
        data_zip_buffer.seek(0)
        st.download_button(
            label="📥 데이터 ZIP 파일 다운로드",
            data=data_zip_buffer,
            file_name=f"{project_title}_데이터_아카이브_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            mime="application/zip",
            key="download_data_zip_actual", # 키 변경
            help="저장된 데이터 ZIP 파일을 다운로드합니다."
        )
        st.success("데이터가 성공적으로 ZIP 파일로 저장되었습니다. 다운로드 버튼을 클릭하세요.")
    else:
        st.warning("⚠️ 저장할 데이터가 없습니다. 요구사항, 일정 또는 변경 이력을 먼저 입력해주세요.")

st.markdown("##### 6.2. 저장된 데이터 불러오기")
st.info("💡 이전에 저장한 데이터 ZIP 파일을 업로드하여 요구사항, 일정 및 문서 이력 데이터를 복원합니다. 현재 입력된 데이터는 덮어쓰여집니다.")

uploaded_data_zip = st.file_uploader(
    "저장된 데이터 ZIP 파일 업로드 (requirements_data.csv, schedule_data.csv, document_history_data.csv 포함)",
    type=["zip"],
    key="upload_data_zip",
    help="요구사항, 일정 및 문서 이력 데이터를 포함하는 ZIP 파일을 업로드하세요."
)

if uploaded_data_zip is not None:
    try:
        with zipfile.ZipFile(uploaded_data_zip, 'r') as zf:
            req_found = False
            schedule_found = False
            history_found = False

            for name in zf.namelist():
                if name == "requirements_data.csv":
                    with zf.open(name) as file:
                        st.session_state.requirements = pd.read_csv(io.TextIOWrapper(file, 'utf-8-sig'))
                        req_found = True
                elif name == "schedule_data.csv":
                    with zf.open(name) as file:
                        # Coordinator 컬럼을 포함하여 읽기
                        temp_schedule_df = pd.read_csv(io.TextIOWrapper(file, 'utf-8-sig'))
                        if 'Coordinator' not in temp_schedule_df.columns:
                            temp_schedule_df['Coordinator'] = None # 없으면 추가
                        st.session_state.schedule_data = temp_schedule_df
                        schedule_found = True
                elif name == "document_history_data.csv":
                    with zf.open(name) as file:
                        st.session_state.document_history = pd.read_csv(io.TextIOWrapper(file, 'utf-8-sig'))
                        history_found = True
            
            if req_found or schedule_found or history_found:
                st.success("✅ 데이터가 성공적으로 불러와졌습니다.")
                # ID 카운터 복원 (가장 큰 ID + 1)
                if not st.session_state.requirements.empty:
                    functional_ids = st.session_state.requirements[st.session_state.requirements['요구사항ID'].str.startswith('RQ-F-', na=False)]['요구사항ID'].str.extract(r'RQ-F-(\d+)').astype(float)  # float으로 변환해서 NaN 처리 용이하도록
                    nonfunctional_ids = st.session_state.requirements[st.session_state.requirements['요구사항ID'].str.startswith('RQ-NF-', na=False)]['요구사항ID'].str.extract(r'RQ-NF-(\d+)').astype(float)

                # 최대값이 존재하는지 검사 후 int 변환
                if not functional_ids.empty and pd.notna(functional_ids.max().squeeze()):
                    next_id_f = int(functional_ids.max().squeeze()) + 1
                else:
                    next_id_f = 1

                if not nonfunctional_ids.empty and pd.notna(nonfunctional_ids.max().squeeze()):
                    next_id_nf = int(nonfunctional_ids.max().squeeze()) + 1
                else:
                    next_id_nf = 1

                st.session_state.next_id_f = next_id_f
                st.session_state.next_id_nf = next_id_nf

                st.info(f"ID 카운터가 복원되었습니다. 다음 기능 요구사항 ID: RQ-F-{next_id_f:03d}, 다음 비기능 요구사항 ID: RQ-NF-{next_id_nf:03d}")

            else:
                st.session_state.next_id_f = 1
                st.session_state.next_id_nf = 1
                st.info("불러온 요구사항이 없어 ID 카운터가 1로 초기화되었습니다.")


                # 데이터 불러오기 후 UI 업데이트를 위해 rerun
                st.rerun() # Use st.rerun() to force an immediate refresh
    except zipfile.BadZipFile:
        st.error("⚠️ 올바른 ZIP 파일이 아닙니다. 손상되었거나 지원되지 않는 형식일 수 있습니다.")
    except Exception as e:
        st.error(f"데이터 불러오기 중 오류가 발생했습니다: {e}")